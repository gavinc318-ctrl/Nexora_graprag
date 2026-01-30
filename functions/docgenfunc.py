from __future__ import annotations

import hashlib
import json
import re
from pathlib import Path
from typing import Any, Dict, Iterable, List, Optional, Sequence, Tuple, Union

from docx import Document

from functions.docgen_schema import DocGenPayload, validate_docgen_payload

_PLACEHOLDER_RE = re.compile(r"\{\{([A-Z0-9_]+)\}\}")


def _hash_bytes(data: bytes) -> str:
    return hashlib.sha256(data).hexdigest()


def _iter_paragraphs(doc: Document) -> Iterable:
    for p in doc.paragraphs:
        yield p
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    yield p


def _extract_placeholders(doc: Document) -> List[str]:
    found: List[str] = []
    for p in _iter_paragraphs(doc):
        txt = "".join(r.text for r in p.runs)
        for m in _PLACEHOLDER_RE.finditer(txt):
            found.append(m.group(1))
    return sorted(set(found))


def _replace_in_text(text: str, mapping: Dict[str, str]) -> str:
    out = text
    for key, val in mapping.items():
        out = out.replace(f"{{{{{key}}}}}", val)
    return out


def _contains_placeholders(text: str) -> bool:
    return bool(_PLACEHOLDER_RE.search(text or ""))


def _replace_in_paragraph(paragraph, mapping: Dict[str, str]) -> None:
    if not paragraph.runs:
        return
    full = "".join(r.text for r in paragraph.runs)
    if not full or not _contains_placeholders(full):
        return

    # First, replace within each run (keeps run styles intact if placeholders stay in one run).
    for run in paragraph.runs:
        if run.text and _contains_placeholders(run.text):
            run.text = _replace_in_text(run.text, mapping)

    full_after = "".join(r.text for r in paragraph.runs)
    if _contains_placeholders(full_after):
        replaced = _replace_in_text(full, mapping)
        paragraph.runs[0].text = replaced
        for run in paragraph.runs[1:]:
            run.text = ""


def _render_evidence_notes(notes: Sequence) -> str:
    if not notes:
        return ""
    lines: List[str] = []
    for i, note in enumerate(notes, start=1):
        topic = note.topic
        lines.append(f"备注{i}：{topic}")
        for j, item in enumerate(note.items or [], start=1):
            lines.append(f"{j}. 关于“{item.claim}”的依据：")
            for ev in item.evidence or []:
                lines.append(f"   - 来源：{ev.source}，第 {ev.page} 页")
                lines.append(f"     摘录：{ev.excerpt}")
        if i != len(notes):
            lines.append("")
    return "\n".join(lines).strip()


def _build_placeholder_mapping(payload: DocGenPayload, template_keys: Sequence[str]) -> Dict[str, str]:
    mapping: Dict[str, str] = {}
    if payload.title:
        mapping["TITLE"] = payload.title
    if payload.doc_no:
        mapping["DOC_NO"] = payload.doc_no

    for sec in payload.sections or []:
        mapping[sec.key] = sec.text

    evidence_text = _render_evidence_notes(payload.evidence_notes or [])
    if evidence_text:
        mapping["EVIDENCE_NOTES"] = evidence_text

    for key in template_keys:
        if key not in mapping:
            mapping[key] = "（待补充）"
    return mapping


def _write_meta(
    meta_path: str,
    payload: DocGenPayload,
    template_path: str,
    template_sha256: str,
    placeholders: Sequence[str],
) -> None:
    payload_json = payload.model_dump(mode="json")
    payload_bytes = json.dumps(payload_json, sort_keys=True, ensure_ascii=True).encode("utf-8")
    meta = {
        "title": payload.title,
        "doc_no": payload.doc_no,
        "sections": payload_json.get("sections", []),
        "evidence_notes": payload_json.get("evidence_notes", []),
        "template_path": str(Path(template_path).name),
        "template_sha256": template_sha256,
        "payload_sha256": _hash_bytes(payload_bytes),
        "placeholders": list(placeholders),
    }
    Path(meta_path).write_text(json.dumps(meta, ensure_ascii=True, sort_keys=True, indent=2), encoding="utf-8")


def load_payload(data: Union[str, Dict[str, Any]]) -> DocGenPayload:
    if isinstance(data, str):
        try:
            raw = json.loads(data)
        except json.JSONDecodeError as e:
            raise ValueError(f"invalid JSON: {e}") from e
    elif isinstance(data, dict):
        raw = data
    else:
        raise TypeError("payload must be JSON string or dict")
    if not isinstance(raw, dict):
        raise ValueError("payload must be a JSON object")
    return validate_docgen_payload(raw)


def render_docx(
    template_path: str,
    output_path: str,
    payload: Union[str, Dict[str, Any], DocGenPayload],
    meta_path: Optional[str] = None,
    create_pdf: bool = False,
    pdf_path: Optional[str] = None,
) -> Dict[str, Any]:
    template = Path(template_path)
    if not template.exists():
        raise FileNotFoundError(f"template not found: {template_path}")

    if isinstance(payload, DocGenPayload):
        data = payload
    else:
        data = load_payload(payload)

    doc = Document(str(template))
    placeholders = _extract_placeholders(doc)
    mapping = _build_placeholder_mapping(data, placeholders)

    for p in _iter_paragraphs(doc):
        _replace_in_paragraph(p, mapping)

    out = Path(output_path)
    out.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out))

    tpl_hash = _hash_bytes(template.read_bytes())
    if meta_path:
        _write_meta(meta_path, data, template_path, tpl_hash, placeholders)

    if create_pdf:
        if not pdf_path:
            pdf_path = str(out.with_suffix(".pdf"))
        _convert_to_pdf(str(out), pdf_path)

    return {
        "ok": True,
        "docx_path": str(out),
        "meta_path": str(meta_path) if meta_path else "",
        "pdf_path": str(pdf_path) if create_pdf else "",
        "placeholders": placeholders,
    }


def _convert_to_pdf(docx_path: str, pdf_path: str) -> None:
    try:
        from docx2pdf import convert  # type: ignore
    except Exception as e:  # pragma: no cover - optional dependency
        raise RuntimeError("docx2pdf not available for PDF export") from e
    convert(docx_path, pdf_path)
