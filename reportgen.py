"""Gradio UI for DocGen (public document generation)."""

from __future__ import annotations

import json
import uuid
from pathlib import Path
import re
from typing import Any, Dict, Optional, Tuple

import gradio as gr

from functions.docgenfunc import render_docx, load_payload as _load_payload_model

_PLACEHOLDER_RE = re.compile(r"\{\{([A-Z0-9_]+)\}\}")

_SAMPLE_JSON = {
    "title": "关于开展专项检查的通知",
    "doc_no": "〔2025〕12号",
    "sections": [
        {"key": "SECTION_1_INTRO", "text": "为进一步规范相关工作流程，现就专项检查事项通知如下。"},
        {"key": "SECTION_2_FACTS", "text": "经初步核查，相关单位已建立基础台账，但执行一致性有待加强。"},
        {"key": "SECTION_3_ANALYSIS", "text": "从制度执行与风险防控角度看，需完善监督机制并强化闭环管理。"},
        {"key": "SECTION_4_DECISION", "text": "决定自本通知之日起开展专项检查，形成整改闭环。"},
    ],
    "evidence_notes": [
        {
            "topic": "基本情况依据",
            "items": [
                {
                    "claim": "已建立基础台账",
                    "evidence": [
                        {"source": "工作报告A.pdf", "page": 3, "excerpt": "各单位已完成台账登记。"}
                    ],
                }
            ],
        }
    ],
}


def _load_payload_text(payload_text: str, payload_file) -> str:
    if payload_file is not None and getattr(payload_file, "name", ""):
        return Path(payload_file.name).read_text(encoding="utf-8").strip()
    return (payload_text or "").strip()


def _safe_output_name(name: str) -> str:
    name = (name or "").strip()
    if not name:
        return "docgen_output.docx"
    if not name.lower().endswith(".docx"):
        name += ".docx"
    return name


def _generate_doc(
    template_file,
    payload_text: str,
    payload_file,
    output_name: str,
    create_pdf: bool,
) -> Tuple[str, Optional[str], Optional[str], Optional[str]]:
    if template_file is None or not getattr(template_file, "name", ""):
        return "❌ 请先上传模板 .docx 文件", None, None, None

    payload_raw = _load_payload_text(payload_text, payload_file)
    if not payload_raw:
        return "❌ 请输入或上传 JSON", None, None, None

    try:
        json.loads(payload_raw)
    except json.JSONDecodeError as e:
        return f"❌ JSON 解析失败：{e}", None, None, None

    out_dir = Path("/tmp/docgen_outputs") / uuid.uuid4().hex
    out_dir.mkdir(parents=True, exist_ok=True)
    output_name = _safe_output_name(output_name)
    docx_path = out_dir / output_name
    meta_path = out_dir / f"{docx_path.stem}_meta.json"
    pdf_path = out_dir / f"{docx_path.stem}.pdf"

    try:
        res = render_docx(
            template_path=template_file.name,
            output_path=str(docx_path),
            payload=payload_raw,
            meta_path=str(meta_path),
            create_pdf=bool(create_pdf),
            pdf_path=str(pdf_path),
        )
    except Exception as e:
        return f"❌ 生成失败：{type(e).__name__}: {e}", None, None, None

    status = "✅ 生成成功"
    docx_out = res.get("docx_path") or str(docx_path)
    meta_out = res.get("meta_path") or str(meta_path)
    pdf_out = res.get("pdf_path") if create_pdf else None
    return status, docx_out, meta_out, pdf_out


def _extract_placeholders_from_template(template_file) -> Tuple[Optional[str], Tuple[str, ...]]:
    if template_file is None or not getattr(template_file, "name", ""):
        return "❌ 请先上传模板 .docx 文件", ()
    from docx import Document  # local import to keep startup light

    doc = Document(template_file.name)
    found = set()
    for p in doc.paragraphs:
        txt = "".join(r.text for r in p.runs)
        for m in _PLACEHOLDER_RE.finditer(txt):
            found.add(m.group(1))
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    txt = "".join(r.text for r in p.runs)
                    for m in _PLACEHOLDER_RE.finditer(txt):
                        found.add(m.group(1))
    return None, tuple(sorted(found))


def _validate_payload(
    template_file,
    payload_text: str,
    payload_file,
) -> str:
    err, placeholders = _extract_placeholders_from_template(template_file)
    if err:
        return err

    payload_raw = _load_payload_text(payload_text, payload_file)
    if not payload_raw:
        return "❌ 请输入或上传 JSON"

    try:
        payload = _load_payload_model(payload_raw)
    except Exception as e:
        return f"❌ JSON 校验失败：{type(e).__name__}: {e}"

    keys = {"TITLE", "DOC_NO", "EVIDENCE_NOTES"}
    for sec in payload.sections or []:
        keys.add(sec.key)

    missing = [k for k in placeholders if k not in keys]
    extra = [k for k in keys if k not in placeholders]

    lines = ["✅ 模板占位符校验通过"]
    if missing:
        lines.append("⚠️ 模板中存在未提供内容的占位符：")
        lines.extend([f"- {k}" for k in missing])
    if extra:
        lines.append("⚠️ JSON 中存在模板未使用的占位符：")
        lines.extend([f"- {k}" for k in extra])
    if not missing and not extra:
        lines.append("占位符与 JSON 对齐，无缺失项。")
    return "\n".join(lines)


def build_app() -> gr.Blocks:
    with gr.Blocks(title="DocGen Report") as demo:
        gr.Markdown("# DocGen 公文生成")
        gr.Markdown("上传模板、填写结构化 JSON，生成固定格式的公文输出。")

        with gr.Row():
            template_file = gr.File(
                label="Word 模板（.docx）",
                file_types=[".docx"],
            )
            payload_file = gr.File(
                label="JSON 文件（可选）",
                file_types=[".json"],
            )

        payload_text = gr.Code(
            label="LLM JSON 输出",
            value=json.dumps(_SAMPLE_JSON, ensure_ascii=True, indent=2),
            language="json",
            lines=18,
        )

        with gr.Row():
            output_name = gr.Textbox(label="输出文件名", value="docgen_output.docx")
        create_pdf = gr.Checkbox(label="同时导出 PDF（需要 docx2pdf）", value=False)

        with gr.Row():
            generate_btn = gr.Button("生成公文")
            validate_btn = gr.Button("校验占位符")

        with gr.Row():
            status = gr.Textbox(label="状态", lines=2, interactive=False)
            out_docx = gr.File(label="输出 DOCX")
            out_meta = gr.File(label="meta.json")
            out_pdf = gr.File(label="输出 PDF（可选）")

        validate_result = gr.Textbox(label="占位符校验结果", lines=8, interactive=False)

        generate_btn.click(
            fn=_generate_doc,
            inputs=[template_file, payload_text, payload_file, output_name, create_pdf],
            outputs=[status, out_docx, out_meta, out_pdf],
        )
        validate_btn.click(
            fn=_validate_payload,
            inputs=[template_file, payload_text, payload_file],
            outputs=[validate_result],
        )

    return demo


if __name__ == "__main__":
    build_app().launch()
